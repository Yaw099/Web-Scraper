from pathlib import Path
from urllib.parse import urlparse, unquote
import requests
import pandas as pd
from pypdf import PdfReader
from docx import Document
import shutil
import subprocess
from bs4 import BeautifulSoup
from urllib.parse import urljoin


DOCUMENT_EXTENSIONS = {
    ".pdf", ".doc", ".docx",
    ".txt", ".xls", ".xlsx",
    ".csv", ".rtf"
}

def discover_document_links(html, base_url):
    soup = BeautifulSoup(html, "lxml")
    document_urls = []

    for tag in soup.find_all("a", href=True):
        absolute_url = urljoin(base_url, tag["href"])

        if is_document_url(absolute_url):
            document_urls.append(absolute_url)

    return list(dict.fromkeys(document_urls))

def find_libreoffice():
    # First check PATH
    path = shutil.which("soffice") or shutil.which("libreoffice")
    if path:
        return path

    # Common Windows install locations
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]

    for candidate in candidates:
        if Path(candidate).exists():
            return candidate

    return None

def normalize_document(path):
    path = Path(path)
    ext = path.suffix.lower()

    conversions = {
        ".doc": "docx",
        ".xls": "xlsx",
        ".ppt": "pptx",
    }

    if ext not in conversions:
        return path

    libreoffice_path = find_libreoffice()

    if not libreoffice_path:
        raise RuntimeError(
            f"Legacy Office file detected: {path.name}. "
            "LibreOffice is required to convert this file."
        )

    output_format = conversions[ext]
    output_folder = path.parent

    subprocess.run(
        [
            libreoffice_path,
            "--headless",
            "--convert-to",
            output_format,
            "--outdir",
            str(output_folder),
            str(path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    converted_path = path.with_suffix(f".{output_format}")

    if not converted_path.exists():
        raise FileNotFoundError(f"Converted file not found: {converted_path}")
    
    # Conversion succeeded, remove the legacy file.
    try:
        path.unlink()
    except Exception:
        # Not critical if cleanup fails.
        pass

    return converted_path

def is_document_url(url: str) -> bool:
    path = urlparse(url).path.lower()
    return Path(path).suffix in DOCUMENT_EXTENSIONS


def download_document(url, output_folder):
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    parsed = urlparse(url)
    filename = Path(unquote(parsed.path)).name

    if not filename:
        filename = "downloaded_document"

    local_path = output_folder / filename

    # Skip download if we already have the original
    if local_path.exists():
        return local_path

    # Skip download if a converted Office document already exists
    suffix = local_path.suffix.lower()

    if suffix == ".doc":
        converted = local_path.with_suffix(".docx")
    elif suffix == ".xls":
        converted = local_path.with_suffix(".xlsx")
    elif suffix == ".ppt":
        converted = local_path.with_suffix(".pptx")
    else:
        converted = None

    if converted and converted.exists():
        return converted
    
    response = requests.get(url, stream=True, timeout=30)
    response.raise_for_status()

    with open(local_path, "wb") as file:
        for chunk in response.iter_content(chunk_size=8192):
            if chunk:
                file.write(chunk)

    return local_path


def get_document_type(path):
    return Path(path).suffix.lower()


def extract_txt(path):
    path = Path(path)

    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="latin-1")


def extract_pdf(path):
    reader = PdfReader(path)

    pages_text = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text()

        if text:
            pages_text.append(f"\n--- Page {page_number} ---\n{text}")

    return "\n".join(pages_text)


def extract_docx(path):
    doc = Document(path)

    output = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            output.append(text)

    for table_index, table in enumerate(doc.tables, start=1):
        output.append(f"\n--- Table {table_index} ---")

        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                cells.append(cell_text)

            # Avoid duplicate repeated cells caused by merged Word cells
            cleaned_cells = []
            for cell in cells:
                if not cleaned_cells or cleaned_cells[-1] != cell:
                    cleaned_cells.append(cell)

            output.append(" | ".join(cleaned_cells))

    return "\n".join(output)


def extract_excel(path):
    sheets = pd.read_excel(path, sheet_name=None)

    output = []

    for sheet_name, df in sheets.items():
        output.append(f"\n--- Sheet: {sheet_name} ---\n")
        output.append(df.to_string(index=False))

    return "\n".join(output)


def extract_csv(path):
    df = pd.read_csv(path)
    return df.to_string(index=False)


def extract_text(path):
    ext = get_document_type(path)

    if ext == ".pdf":
        return extract_pdf(path)

    elif ext == ".docx":
        return extract_docx(path)

    elif ext == ".txt":
        return extract_txt(path)

    elif ext in [".xls", ".xlsx"]:
        return extract_excel(path)

    elif ext == ".csv":
        return extract_csv(path)
    else:
        raise ValueError(f"Unsupported document type: {ext}")
    

def process_documents_from_links(links, output_folder):
    results = []

    for link in links:
        url = link.get("url") if isinstance(link, dict) else link

        if not url:
            continue

        if not is_document_url(url):
            continue

        result = process_document(url, output_folder)
        results.append(result)

    return results

def process_document(url, output_folder):
    try:
        local_file = download_document(url, output_folder)
        normalized_file = normalize_document(local_file)
        text = extract_text(normalized_file)

        return {
            "url": url,
            "local_path": str(local_file),
            "normalized_path": str(normalized_file),
            "file_type": get_document_type(normalized_file),
            "text": text,
            "success": True,
            "error": None
        }

    except Exception as error:
        return {
            "url": url,
            "local_path": None,
            "normalized_path": None,
            "file_type": None,
            "text": "",
            "success": False,
            "error": str(error)
        }

def process_downloaded_document_file(path):
    normalized_file = normalize_document(path)
    text = extract_text(normalized_file)

    return {
        "local_path": str(path),
        "normalized_path": str(normalized_file),
        "file_type": get_document_type(normalized_file),
        "text": text,
        "success": True,
        "error": None,
    }

def download_documents_from_links(links, output_folder):
    results = []

    for link in links:
        url = link.get("url") if isinstance(link, dict) else link

        if not url or not is_document_url(url):
            continue

        try:
            local_file = download_document(url, output_folder)

            results.append({
                "url": url,
                "local_path": str(local_file),
                "normalized_path": "",
                "file_type": get_document_type(local_file),
                "text": "",
                "success": True,
                "error": None,
            })

        except Exception as error:
            results.append({
                "url": url,
                "local_path": "",
                "normalized_path": "",
                "file_type": "",
                "text": "",
                "success": False,
                "error": str(error),
            })

    return results

# extract_tables_pdf()
# extract_sheet_names()
# extract_headings()
# extract_pptx()